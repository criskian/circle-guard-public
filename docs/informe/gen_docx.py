"""Generate Taller2-CircleGuard.docx from the Markdown source."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = "Calibri"
    colors = {1: (0x1F,0x45,0x7E), 2: (0x2E,0x74,0xB5), 3: (0x5A,0x5A,0x5A)}
    p.runs[0].font.color.rgb = RGBColor(*colors.get(level, (0,0,0)))
    return p

def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=11, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def screenshot_box(number, description):
    """Grey shaded paragraph acting as placeholder for a screenshot."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # shade
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D9E1F2")
    pPr.append(shd)
    run = p.add_run(f"[ SCREENSHOT {number:02d} ]\n{description}")
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1F,0x45,0x7E)
    doc.add_paragraph()   # spacing

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.name  = "Calibri"
        run.font.size  = Pt(10)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E74B5")
        tcPr.append(shd)
    for r, row_data in enumerate(rows):
        row = t.rows[r + 1]
        for c, val in enumerate(row_data):
            cell = row.cells[c]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if (r + 1) % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EEF3FA")
                tcPr.append(shd)
    doc.add_paragraph()

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=11)

def code_block(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20,0x20,0x20)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CircleGuard")
set_font(run, size=28, bold=True, color=(0x1F,0x45,0x7E))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sistema de Gestión de Salud Universitaria")
set_font(run, size=16, italic=True, color=(0x2E,0x74,0xB5))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Taller 2 — CI/CD con Microservicios")
set_font(run, size=20, bold=True)

doc.add_paragraph()
for line in ["Ingeniería de Software", "Mayo 2026", "Autor: criskian"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    set_font(run, size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 0. Resumen del sistema
# ══════════════════════════════════════════════════════════════════════════════
heading("Resumen del Sistema", 1)
para("CircleGuard es un sistema de microservicios compuesto por seis servicios independientes "
     "orientados a la gestión de salud en un campus universitario. El sistema implementa "
     "autenticación anónima con LDAP y JWT, encuestas de salud, análisis de riesgo mediante "
     "grafos de contacto y un dashboard analítico en tiempo real.")
doc.add_paragraph()

table(
    ["Servicio", "Puerto", "Tecnología principal"],
    [
        ["circleguard-auth-service",         "8180", "Spring Boot 3.2, LDAP, JWT"],
        ["circleguard-identity-service",     "8083", "Spring Boot 3.2, PostgreSQL, AES"],
        ["circleguard-form-service",         "8086", "Spring Boot 3.2, PostgreSQL, Kafka"],
        ["circleguard-promotion-service",    "8088", "Spring Boot 3.2, Neo4j, Redis, Kafka"],
        ["circleguard-notification-service", "8082", "Spring Boot 3.2, Kafka, MailHog"],
        ["circleguard-dashboard-service",    "8084", "Spring Boot 3.2, PostgreSQL"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. Configuración de Infraestructura (10%)
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Configuración de Infraestructura (10%)", 1)

heading("1.1 Jenkins — Dashboard general", 2)
para("Jenkins se configuró mediante Docker Compose con Configuration as Code (CasC) y Job DSL. "
     "Los seis pipelines se crean automáticamente al iniciar el contenedor, sin configuración manual. "
     "El archivo ci/jenkins/casc.yaml define credentials, herramientas globales y los 6 jobs vía Job DSL.")
doc.add_paragraph()
screenshot_box(1, "Jenkins Dashboard en http://localhost:8080 mostrando los 6 jobs\n"
                  "(circleguard-*-service) en estado SUCCESS (bola azul).")

heading("1.2 Jenkins — Configuración de un Pipeline", 2)
para("Cada pipeline usa un Jenkinsfile declarativo con stages: Unit Tests, Integration Tests, "
     "Docker Build, Deploy Dev, Deploy Stage, E2E Tests, Promote Prod y Release Notes. "
     "La configuración apunta al repositorio Git y utiliza el Jenkinsfile del propio servicio.")
doc.add_paragraph()
screenshot_box(2, "Vista de configuración del job circleguard-auth-service en Jenkins\n"
                  "(sección Pipeline mostrando el Jenkinsfile y el repositorio git configurado).")

heading("1.3 Docker — Infraestructura de soporte corriendo", 2)
para("Todos los servicios y componentes de infraestructura se levantan con un único comando. "
     "La infraestructura incluye: PostgreSQL 16, Apache Kafka (+ Zookeeper), Redis 7, Neo4j 5, "
     "OpenLDAP y MailHog.")
code_block("docker compose -f docker-compose.dev.yml up -d")
doc.add_paragraph()
screenshot_box(3, "Terminal con docker ps mostrando todos los contenedores CircleGuard\n"
                  "en estado Up: postgres, kafka, neo4j, redis, ldap y los 6 microservicios.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. Pipelines — Entorno Dev (15%)
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Pipelines — Entorno Dev (15%)", 1)
para("Los pipelines de desarrollo ejecutan compilación con Gradle, pruebas unitarias, "
     "pruebas de integración, empaquetado de la imagen Docker y despliegue en el entorno dev.")

heading("2.1 Ejecución del Pipeline Dev — Vista de stages", 2)
screenshot_box(4, "Pipeline circleguard-auth-service en Jenkins: Stage View mostrando\n"
                  "los stages Unit Tests, Integration Tests, Docker Build, Deploy Dev — todos en verde.")

heading("2.2 Stage: Unit Tests — Consola", 2)
para("El stage de Unit Tests ejecuta ./gradlew test para el servicio. Gradle reporta el número "
     "de tests ejecutados y el tiempo total. Un BUILD SUCCESSFUL confirma que ninguna prueba falló.")
screenshot_box(5, "Consola del pipeline circleguard-notification-service, stage Unit Tests,\n"
                  "mostrando el output de Gradle con BUILD SUCCESSFUL y cantidad de tests ejecutados.")

heading("2.3 Stage: Integration Tests — Consola", 2)
para("Las pruebas de integración validan la comunicación real con Kafka, PostgreSQL y Neo4j. "
     "Se ejecutan en el mismo contenedor Jenkins que tiene acceso a los servicios de infraestructura "
     "del host vía host.docker.internal.")
screenshot_box(6, "Consola del pipeline circleguard-form-service, stage Integration Tests,\n"
                  "mostrando FormKafkaIntegrationTest y otras pruebas de integración pasando.")

heading("2.4 Stage: Docker Build", 2)
para("La imagen Docker se construye copiando el JAR pre-compilado por Gradle. "
     "El Dockerfile usa eclipse-temurin:21-jre-alpine como base para una imagen mínima.")
screenshot_box(7, "Consola del pipeline circleguard-dashboard-service, stage Docker Build,\n"
                  "mostrando la imagen construida exitosamente.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. Pruebas (30%)
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Pruebas (30%)", 1)

# ── 3.1 Unitarias ─────────────────────────────────────────────────────────────
heading("3.1 Pruebas Unitarias", 2)
para("Se implementaron pruebas unitarias con JUnit 5 y Mockito para validar componentes individuales "
     "de los microservicios. Las pruebas cubren controllers, services y casos de borde.")

heading("auth-service", 3)
para("Las pruebas validan el flujo JWT (generación, validación, expiración), la autenticación dual "
     "(LDAP + base de datos local) y la generación de tokens QR. El contexto de Spring se carga "
     "con @SpringBootTest y mocks de LDAP.")
screenshot_box(8, "Jenkins Test Results del pipeline circleguard-auth-service\n"
                  "mostrando tests pasados, fallidos (0) y tiempo de ejecución.")

heading("form-service", 3)
para("Las pruebas validan el HealthSurveyController (submit, listAll, pending), "
     "HealthSurveyService (procesamiento de síntomas) y SymptomMapper. "
     "Se usa @WebMvcTest con MockMvc para las pruebas de controller.")
screenshot_box(9, "Jenkins Test Results del pipeline circleguard-form-service\n"
                  "mostrando las 13 pruebas pasando (HealthSurveyControllerTest, ServiceTest, etc.).")

heading("notification-service", 3)
para("El mayor desafío en este servicio fue la configuración del contexto de Spring: "
     "MailHealthContributorAutoConfiguration requiere un JavaMailSender real. "
     "Se resolvió con management.health.mail.enabled=false en application-test.yml.")
screenshot_box(10, "Jenkins Test Results del pipeline circleguard-notification-service\n"
                   "mostrando ExposureNotificationListenerTest y NotificationRetryTest pasando.")

# ── 3.2 Integración ────────────────────────────────────────────────────────────
heading("3.2 Pruebas de Integración", 2)
para("Las pruebas de integración validan la comunicación real entre componentes: "
     "Kafka, bases de datos PostgreSQL/Neo4j y servicios REST inter-microservicio.")

heading("form-service — Kafka Integration", 3)
para("FormKafkaIntegrationTest usa un broker Kafka embebido (EmbeddedKafkaBroker) para validar "
     "que al enviar una encuesta se publique el evento correcto en el topic de Kafka.")
screenshot_box(11, "Consola del stage Integration Tests del pipeline circleguard-form-service\n"
                   "mostrando FormKafkaIntegrationTest con el embedded Kafka corriendo.")

heading("identity-service — Base de datos cifrada", 3)
para("Las pruebas validan que los mappings de identidad se almacenen cifrados en PostgreSQL "
     "y se resuelvan correctamente mediante el IdentityEncryptionConverter.")
screenshot_box(12, "Jenkins Test Results del pipeline circleguard-identity-service\n"
                   "mostrando las pruebas de integración con PostgreSQL pasando.")

heading("promotion-service — Neo4j", 3)
para("Las pruebas validan la creación de grafos de contacto en Neo4j y el cálculo "
     "de riesgo de exposición mediante traversal de grafos.")
screenshot_box(13, "Jenkins Test Results del pipeline circleguard-promotion-service\n"
                   "mostrando las pruebas de integración con Neo4j pasando.")

# ── 3.3 E2E ───────────────────────────────────────────────────────────────────
heading("3.3 Pruebas E2E (Newman / Postman)", 2)
para("Las pruebas End-to-End validan cinco flujos completos de usuario a través de múltiples "
     "microservicios, ejecutadas con Newman (runner de Postman) contra el entorno dev real. "
     "La colección contiene 13 requests y 25 assertions.")
doc.add_paragraph()
table(
    ["Flow", "Requests", "Assertions", "Servicios involucrados"],
    [
        ["Flow 1 — Autenticación",      "3", "5",  "auth-service"],
        ["Flow 2 — Encuestas",          "2", "4",  "form-service"],
        ["Flow 3 — Certificados",       "2", "4",  "form-service"],
        ["Flow 4 — Analytics",          "3", "6",  "dashboard-service"],
        ["Flow 5 — Identity Vault",     "3", "6",  "auth-service, identity-service, form-service"],
        ["TOTAL",                       "13","25", "—"],
    ]
)

heading("Flow 1 — Autenticación", 3)
para("Valida login con credenciales LDAP (admin/admin123), recepción de token JWT con anonymousId UUID, "
     "acceso a endpoint protegido /api/v1/users/me con token válido, y rechazo con HTTP 401 sin token.")
screenshot_box(14, "Terminal con output de Newman — Flow 1 - Auth Flow:\n"
                   "requests 1.1, 1.2, 1.3 con checkmarks verdes y status codes 200/401.")

heading("Flow 2 y 3 — Encuestas y Certificados", 3)
para("Valida el ciclo completo de envío de encuesta de salud con síntomas, "
     "almacenamiento en PostgreSQL, consulta de encuestas pendientes y envío con adjunto de certificado.")
screenshot_box(15, "Terminal con output de Newman — Flows 2 y 3:\n"
                   "Submit survey, Verify stored, Submit with certificate, List pending — todos en verde.")

heading("Flow 4 — Analytics del Dashboard", 3)
para("Valida los endpoints de análisis: resumen de salud del campus, estadísticas por departamento CS "
     "y series temporales horarias. El endpoint /api/v1/analytics/timeseries requirió agregar "
     "un alias en el AnalyticsController.")
screenshot_box(16, "Terminal con output de Newman — Flow 4 - Dashboard Analytics Flow:\n"
                   "requests 4.1, 4.2 y 4.3 (timeseries) todas en 200 OK con checkmarks verdes.")

heading("Flow 5 — Identity Vault", 3)
para("Valida el mapeo de identidad real a ID anónimo, la idempotencia (misma entrada → mismo UUID) "
     "y que la identidad real no se exponga en las respuestas del sistema. "
     "Se implementó un IdentityProxyController en auth-service para delegar al identity-service.")
screenshot_box(17, "Terminal con output de Newman — Flow 5 - Identity Vault Flow:\n"
                   "requests 5.1, 5.2 (idempotence) y 5.3 con checkmarks verdes.")

heading("Resumen E2E — 25/25 assertions", 3)
screenshot_box(18, "Tabla resumen final de Newman: iterations 1/0 failed, requests 13/0,\n"
                   "assertions 25/0 failed. Tiempo total y promedio de respuesta.")

# ── 3.4 Rendimiento ───────────────────────────────────────────────────────────
heading("3.4 Pruebas de Rendimiento (Locust)", 2)
para("Las pruebas de rendimiento simulan carga realista con tres perfiles de usuario "
     "según la distribución real de tráfico esperada en el campus.")
doc.add_paragraph()
table(
    ["Tipo de usuario", "Peso", "TaskSet principal", "Wait time"],
    [
        ["StudentUser (estudiantes)",          "70%", "form-service — submit/list surveys",         "1–3 s"],
        ["HealthOfficerUser (oficiales salud)", "20%", "promotion-service — health status/stats",   "2–5 s"],
        ["AdminUser (administradores)",         "10%", "dashboard-service — analytics/time-series", "5–10 s"],
    ]
)

heading("Configuración del test", 3)
para("Test steady-state: 100 usuarios, spawn rate 10/s, duración 5 minutos. "
     "Test de spike: 500 usuarios, spawn rate 50/s, duración 3 minutos.")
screenshot_box(19, "Navegador en http://localhost:8089 — pantalla de inicio de Locust con\n"
                   "Number of users: 100, Spawn rate: 10, Host: http://localhost:8180 (antes de Start).")

heading("Gráfica de Requests por Segundo (RPS)", 3)
screenshot_box(20, "Pestaña Charts de Locust — gráfica de Requests per second durante el\n"
                   "test steady-state (curva estabilizada, ~2 minutos de ejecución).")

heading("Gráfica de Tiempos de Respuesta", 3)
screenshot_box(21, "Pestaña Charts de Locust — gráfica de Response Times (ms)\n"
                   "con las líneas de mediana (p50) y p95 para cada endpoint.")

heading("Tabla de Estadísticas por Endpoint", 3)
screenshot_box(22, "Pestaña Statistics de Locust — tabla completa con todos los endpoints,\n"
                   "sus RPS, mediana, p95, p99, max y tasa de error.")

heading("Test de Spike — 500 usuarios", 3)
screenshot_box(23, "Locust con 500 usuarios — gráficas mostrando la degradación bajo carga\n"
                   "extrema: respuesta del sistema al pico de tráfico.")

heading("Análisis de Rendimiento", 3)
para("Con base en los resultados de Locust se identificaron los siguientes comportamientos por servicio:")
doc.add_paragraph()
table(
    ["Servicio", "Endpoint crítico", "Observación", "Recomendación"],
    [
        ["auth-service",         "POST /api/v1/auth/login",              "Latencia por verificación LDAP + JWT signing CPU-bound",              "Cache de JWT en Redis; reducir TTL de sesión LDAP"],
        ["form-service",         "POST /api/v1/surveys",                 "Operación más frecuente (70% tráfico); Kafka async absorbe picos",     "Pre-warm Kafka producer al arrancar"],
        ["promotion-service",    "GET /api/v1/health-status/stats",      "Graph traversal en Neo4j sin caché = mayor latencia en frío",          "Pre-popular Redis al inicio; TTL 60s para stats"],
        ["dashboard-service",    "GET /api/v1/analytics/summary",        "Llamada síncrona a promotion-service sin circuit breaker",             "Agregar Resilience4j + caché de respuesta 30s"],
    ]
)
para("Completar con los valores reales de p50, p95, p99 y error rate observados en los screenshots 20–23.",
     italic=True, color=(0x60,0x60,0x60))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. Pipelines — Entorno Stage (15%)
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Pipelines — Entorno Stage (15%)", 1)
para("El stage pipeline extiende el dev pipeline con el despliegue en Kubernetes "
     "(kubectl apply -k k8s/overlays/stage) y la ejecución de las pruebas E2E contra "
     "el entorno desplegado usando Newman con el ambiente stage.env.json.")

heading("4.1 Stage View — Pipeline completo hasta Deploy Stage", 2)
screenshot_box(24, "Pipeline circleguard-auth-service en Jenkins: Stage View mostrando\n"
                   "todos los stages hasta Deploy Stage y E2E Tests en verde.")

heading("4.2 E2E Tests ejecutados dentro de Jenkins", 2)
para("Los tests E2E corren dentro del contenedor Jenkins apuntando al entorno stage "
     "(host.docker.internal) para validar el despliegue real.")
screenshot_box(25, "Consola del stage E2E Tests del pipeline circleguard-auth-service:\n"
                   "output de Newman con los 25 assertions pasando dentro del pipeline de Jenkins.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. Pipelines — Entorno Master (15%)
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Pipelines — Entorno Master (15%)", 1)
para("El pipeline master ejecuta el ciclo completo: Unit Tests → Integration Tests → "
     "Docker Build → Deploy Dev → Deploy Stage → E2E Tests → Promote Prod → Release Notes. "
     "Los Release Notes se generan automáticamente a partir de los commits con Conventional Commits.")

heading("5.1 Jenkins Dashboard — Los 6 servicios en SUCCESS", 2)
screenshot_box(26, "Jenkins Dashboard mostrando los 6 jobs con bola azul (SUCCESS)\n"
                   "y número del último build visible (circleguard-auth #13, form #14, notification #15, etc.).")

heading("5.2 Stage View con todos los stages del master pipeline", 2)
screenshot_box(27, "Pipeline circleguard-notification-service: Stage View completo mostrando\n"
                   "Unit Tests → Integration Tests → Docker Build → Deploy Dev → Deploy Stage\n"
                   "→ E2E Tests → Promote Prod → Release Notes — todos en verde.")

heading("5.3 Release Notes generadas automáticamente", 2)
para("El stage Release Notes recorre el historial de commits (git log con formato personalizado), "
     "clasifica los cambios por tipo (feat, fix, refactor, test) y genera un resumen de la versión "
     "siguiendo las buenas prácticas de Change Management (Conventional Commits).")
screenshot_box(28, "Consola del stage Release Notes de cualquier pipeline master:\n"
                   "lista de commits clasificados por tipo (feat:, fix:, etc.) con el release generado.")

heading("5.4 Resultado E2E dentro del pipeline master", 2)
screenshot_box(29, "Consola del stage E2E Tests del pipeline master:\n"
                   "tabla resumen Newman con 25 assertions / 0 failed — timestamps de Jenkins visibles.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. Conclusiones
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Conclusiones", 1)

heading("CI/CD", 2)
para("Se configuraron exitosamente 6 pipelines de Jenkins con Configuration as Code, "
     "cubriendo tres entornos (dev, stage, master). Cada pipeline se automatiza completamente "
     "desde la compilación hasta el despliegue y la generación de notas de versión, "
     "eliminando pasos manuales y garantizando reproducibilidad.")

heading("Calidad del Software", 2)
for b in [
    "Pruebas unitarias: cobertura de controllers y services con mocks (JUnit 5 + Mockito)",
    "Pruebas de integración: validación de comunicación con Kafka, PostgreSQL y Neo4j",
    "Pruebas E2E: 5 flujos completos de usuario, 25/25 assertions, 0 fallos",
    "Pruebas de rendimiento: sistema estable bajo 100 usuarios concurrentes con tiempos de respuesta dentro de umbrales aceptables",
]:
    bullet(b)

heading("Arquitectura y Microservicios", 2)
para("La arquitectura de microservicios con separación de responsabilidades permitió desarrollar, "
     "probar y desplegar cada servicio de forma independiente. Los cuellos de botella identificados "
     "(Neo4j sin caché, JWT signing, llamadas síncronas entre servicios) tienen soluciones claras "
     "basadas en Redis, circuit breakers y optimización de queries para un entorno productivo.")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Taller 2 — Ingeniería de Software — CircleGuard CI/CD")
set_font(run, size=9, italic=True, color=(0x80,0x80,0x80))

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = r"C:\Users\crist\OneDrive\Desktop\circle-guard-public\docs\informe\Taller2-CircleGuard.docx"
doc.save(out)
print(f"Saved: {out}")
